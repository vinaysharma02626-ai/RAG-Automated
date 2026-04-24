import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


@dataclass
class ParsedBlock:
    text: str
    block_type: str = "text"  # text, heading, list, table
    heading_level: int = 0
    heading_text: Optional[str] = None
    page_num: Optional[int] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    blocks: List[ParsedBlock]
    page_count: int
    title: Optional[str]
    metadata: Dict[str, Any] = field(default_factory=dict)


def parse_pdf(file_path: str) -> ParsedDocument:
    """Extract text and structure from a PDF."""
    import pdfplumber

    blocks: List[ParsedBlock] = []
    page_count = 0
    title = None

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            if pdf.metadata:
                title = pdf.metadata.get("Title") or None

            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract tables first
                tables = page.extract_tables()
                table_bboxes = []
                for table in tables:
                    if table:
                        rows = []
                        for row in table:
                            cleaned = [cell or "" for cell in row]
                            rows.append(" | ".join(cleaned))
                        table_text = "\n".join(rows)
                        blocks.append(
                            ParsedBlock(
                                text=table_text,
                                block_type="table",
                                page_num=page_num,
                            )
                        )

                # Extract text with words for better structure detection
                words = page.extract_words(
                    x_tolerance=3, y_tolerance=3, keep_blank_chars=False
                )
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if not text:
                    continue

                lines = text.split("\n")
                current_para = []
                for line in lines:
                    stripped = line.strip()
                    if not stripped:
                        if current_para:
                            para_text = " ".join(current_para)
                            block_type, heading_level = _classify_line(para_text)
                            blocks.append(
                                ParsedBlock(
                                    text=para_text,
                                    block_type=block_type,
                                    heading_level=heading_level,
                                    page_num=page_num,
                                )
                            )
                            current_para = []
                    else:
                        current_para.append(stripped)

                if current_para:
                    para_text = " ".join(current_para)
                    block_type, heading_level = _classify_line(para_text)
                    blocks.append(
                        ParsedBlock(
                            text=para_text,
                            block_type=block_type,
                            heading_level=heading_level,
                            page_num=page_num,
                        )
                    )

    except Exception as e:
        logger.error(f"PDF parsing error: {e}")
        raise

    # Try to find title from first heading if not in metadata
    if not title:
        for block in blocks[:10]:
            if block.heading_level == 1 or (block.heading_level == 0 and len(block.text) < 100):
                title = block.text[:100]
                break

    return ParsedDocument(blocks=blocks, page_count=page_count, title=title)


def parse_docx(file_path: str) -> ParsedDocument:
    """Extract text and structure from a DOCX file."""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    doc = DocxDocument(file_path)
    blocks: List[ParsedBlock] = []
    title = doc.core_properties.title or None

    heading_map = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Title": 1,
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        heading_level = heading_map.get(style_name, 0)
        block_type = "heading" if heading_level > 0 else "text"

        if style_name and "list" in style_name.lower():
            block_type = "list"

        blocks.append(
            ParsedBlock(
                text=text,
                block_type=block_type,
                heading_level=heading_level,
                page_num=None,
            )
        )

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            blocks.append(ParsedBlock(text=table_text, block_type="table"))

    if not title:
        for block in blocks[:5]:
            if block.heading_level == 1:
                title = block.text[:100]
                break

    return ParsedDocument(blocks=blocks, page_count=0, title=title)


def _classify_line(text: str):
    """Heuristic heading detection for PDFs."""
    if len(text) < 5:
        return "text", 0
    # All caps short line → likely heading
    if text.isupper() and len(text) < 80:
        return "heading", 1
    # Numbered section like 1., 1.1, etc.
    if re.match(r"^\d+(\.\d+)*\.?\s+\w", text) and len(text) < 100:
        dots = text.split(" ")[0].count(".")
        return "heading", min(dots + 1, 3)
    return "text", 0
